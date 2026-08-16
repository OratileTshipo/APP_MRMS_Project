#!/usr/bin/env python3
"""Update APP-MRMS Word docs so list schemas match the CSV source of truth.

CSV-derived ground truth:
  Directorates : DirectorateID(Title), DirectorateName(LinkTitle), Programme(text),
                 Director, DeputyDirector, Deligate (Person), Active (Yes/No)
  Programmes   : Title, ProgrammeCode, ProgrammeManager (Person), Description, Active
  Projects     : Title(ProjectID), ProjectCode, PriorityChallenge, ProjectGoalDescription,
                 KeyDeliverable, AnnualTarget, Reference, Budget (Text), Directorate
                 (Lookup -> Directorates, Title projection), Programme (text),
                 Responsibility (Person), StartDate, CompletionDate, FinancialYear
                 (2026/27..2022/23), Status (Active/Closed/Cancelled)
  Activities   : Title, ActivityCode, ActivityDescription, ActivityShortDescription,
                 Directorate (text), DirectorateLabel, ProgrammeLabel, Project (Lookup),
                 Project: Reference, Frequency, Owner, ActivityOwner, Supervisor,
                 StartDate, EndDate, DueDate, Active, Status (Active/Completed/Not Started)
  MonthlyReports: Title, Programme, Directorate, ProgrammeLabel, DirectorateLabel,
                 Project, ProjectOwnerLabel, Activity (Lookup), ActivityOwnerLabel,
                 Activity: ActivityShortDescription/StartDate/EndDate/DueDate,
                 ReportingMonth (April..March), Quarter (Q1-Q4, All), FinancialYear
                 (2026/27, 2025/26, 2024/25, All), PlannedAct, PercentageComplete,
                 Status (Draft/Submitted/Approved/Rejected/Escalated), Version,
                 SubmittedBy, SubmissionDate, ReviewedBy, ReviewDate, RejectionReason,
                 IsOverdue, isFlagged, ReasonForDeviation, RemedialAction, Output
  APP_Users    : Title, UserAccount (Person), Role (Administrator/ProgrammeManager/
                 Supervisor/Contributor/DeputyDirectorME), Directorate (Lookup ->
                 Directorates, DirectorateID projection), Programme (text), Active
"""
import copy
import docx
from docx.table import _Row


def set_cell(cell, text):
    """Replace a cell's text, preserving the first run's formatting."""
    first_p = cell.paragraphs[0]
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    runs = first_p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        first_p.add_run(text)


def insert_row_after(table, ref_row, texts):
    """Insert a new row (deep copy of ref_row) immediately after ref_row."""
    new_tr = copy.deepcopy(ref_row._tr)
    ref_row._tr.addnext(new_tr)
    idx = list(table._tbl).index(new_tr)
    row = _Row(new_tr, table)
    for cell, text in zip(row.cells, texts):
        set_cell(cell, text)
    return row


def replace_in_paragraph(par, old, new):
    """Replace text in a paragraph across runs (rebuilds run text)."""
    full = "".join(r.text for r in par.runs)
    if old not in full:
        return False
    full = full.replace(old, new)
    for r in par.runs[1:]:
        r._element.getparent().remove(r._element)
    if par.runs:
        par.runs[0].text = full
    else:
        par.add_run(full)
    return True


def replace_in_table(table, find_cell0, new_cells):
    """Update a table row whose first cell contains find_cell0."""
    for row in table.rows:
        if row.cells[0].text.strip().startswith(find_cell0):
            for cell, text in zip(row.cells, new_cells):
                set_cell(cell, text)
            return row
    return None


def replace_in_doc(doc, old, new):
    n = 0
    for par in doc.paragraphs:
        if replace_in_paragraph(par, old, new):
            n += 1
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    if replace_in_paragraph(par, old, new):
                        n += 1
    return n


# ---------------------------------------------------------------- BRD/FDS/TDS
def fix_brd(path):
    d = docx.Document(path)
    changes = []

    # --- 3.2.3 Projects (table 7) ---
    t = d.tables[7]
    r = replace_in_table(t, "Title", ["Title", "Single line text",
        'Project ID (Title display name), e.g. "COMM-PJ-002"; unique & indexed'])
    insert_row_after(t, r, ["ProjectCode", "Single line text", 'Optional project code (e.g. "COMM-PJ-002")'])
    changes.append("Projects: added ProjectCode row")
    replace_in_table(t, "Budget", ["Budget", "Single line text",
        'Rand value; free-text variants stored as text per live CSV schema'])
    changes.append("Projects: Budget type -> Single line text")
    replace_in_table(t, "Directorate", ["Directorate (+ Directorate: Title)", "Lookup (Directorates)",
        'Directorate code; "Directorate: Title" projection shows the full directorate name'])
    changes.append("Projects: Directorate -> Lookup (Directorates)")
    replace_in_table(t, "Programme", ["Programme", "Single line text",
        "⚠ Text, not a Lookup to Programmes — same pattern as Directorates.Programme"])
    changes.append("Projects: Programme -> Single line text")
    replace_in_table(t, "FinancialYear", ["FinancialYear", "Choice", "2026/27 … 2022/23, colour-coded"])
    changes.append("Projects: FinancialYear choices -> 2026/27..2022/23")

    # --- 3.2.4 Activities (table 8) ---
    t = d.tables[8]
    r = replace_in_table(t, "Title", ["Title", "Single line text", 'Activity ID, e.g. "ICTM-AC-0001"'])
    insert_row_after(t, r, ["ActivityCode", "Single line text", 'Activity code (e.g. "ICTM-AC-0001"); also carried in Title'])
    changes.append("Activities: added ActivityCode row")
    r = replace_in_table(t, "ActivityShortDescription", ["ActivityShortDescription", "Single line text", "Denormalized short label for gallery display"])
    insert_row_after(t, r, ["Directorate", "Single line text", 'Directorate short code, e.g. "ICT" (matches Directorates.DirectorateID)'])
    changes.append("Activities: added Directorate row")
    r = replace_in_table(t, "Owner", ["Owner", "Person", ""])
    insert_row_after(t, r, ["ActivityOwner", "Person", "Contributor accountable for capture (separate 'ActivityOwner' column)"])
    changes.append("Activities: split Owner / ActivityOwner")
    replace_in_table(t, "DueDayOfMonth", ["DueDate", "Date only", "Report due date (EndDate + 3 days in current data)"])
    changes.append("Activities: DueDayOfMonth -> DueDate (Date only)")
    r = replace_in_table(t, "Active", ["Active", "Yes/No", "Default = true"])
    insert_row_after(t, r, ["Status", "Choice", "Active / Completed / Not Started, colour-coded"])
    changes.append("Activities: added Status row")

    # --- 3.2.5 MonthlyReports (table 9) ---
    t = d.tables[9]
    replace_in_table(t, "Activity / ActivityOwnerLabel", ["Activity / ActivityOwnerLabel", "Lookup / Text",
        "Plus read-only Activity: ActivityShortDescription / StartDate / EndDate / DueDate projections"])
    changes.append("MonthlyReports: Activity projections DueDayOfMonth -> DueDate")
    replace_in_table(t, "Quarter", ["Quarter", "Choice", "Q1–Q4 (+ All)"])
    changes.append("MonthlyReports: Quarter choices -> Q1-Q4 (+ All)")
    r = replace_in_table(t, "ProgressNarrative", ["ReasonForDeviation", "Multiple lines of text", "Free-text reason when progress deviates from plan"])
    r = insert_row_after(t, r, ["RemedialAction", "Multiple lines of text", "Corrective action taken/planned"])
    insert_row_after(t, r, ["Output", "Multiple lines of text", "Reported output / achievement narrative"])
    changes.append("MonthlyReports: ProgressNarrative -> ReasonForDeviation, RemedialAction, Output")
    replace_in_table(t, "Status", ["Status", "Choice", "Draft / Submitted / Approved / Rejected / Escalated"])
    changes.append("MonthlyReports: Status choices include Escalated")

    # --- prose/code fixes ---
    n = replace_in_doc(d, "ProgressNarrative", "Output")
    changes.append(f"BRD prose: ProgressNarrative -> Output ({n} replacements)")
    n = replace_in_doc(d, "DueDayOfMonth", "DueDate")
    changes.append(f"BRD prose: DueDayOfMonth -> DueDate ({n} replacements)")
    n = replace_in_doc(d,
        "Projects.csv and Projects_1.csv exported as identical duplicates — confirm with you whether this is one physical list surfaced twice, or two lists that have since diverged, before any script references \"Projects_1\".",
        "⚠ The live CSV export contains a single Projects list (Projects.csv). There is no separate \"Projects_1\" list — the earlier duplicate export was a snapshot artefact and should not be referenced.")
    changes.append(f"BRD prose: removed stale Projects_1 note ({n} replacements)")

    d.save(path)
    return changes


# ------------------------------------------------------------ Architecture Pack
def fix_arch(path):
    d = docx.Document(path)
    changes = []

    # --- 4.1 Programmes (table 10) ---
    t = d.tables[10]
    replace_in_table(t, "Title", ["Title", "Single line text",
        "Programme name, e.g. 'Programme 1', 'Programme 2A', 'Programme 2B', 'Programme 3', 'Programme 4'"])
    changes.append("Programmes: corrected example programme names")

    # --- 4.2 Projects (table 11) ---
    t = d.tables[11]
    r = replace_in_table(t, "Title", ["Title", "Single line text",
        "Project ID (Title display name), e.g. 'COMM-PJ-002'; unique & indexed"])
    insert_row_after(t, r, ["ProjectCode", "Single line text", "Optional project code (e.g. 'COMM-PJ-002')"])
    r = replace_in_table(t, "PriorityChallenge", ["PriorityChallenge", "Multiple lines of text", ""])
    insert_row_after(t, r, ["ProjectGoalDescription", "Multiple lines of text", "Project goal description"])
    replace_in_table(t, "Budget", ["Budget", "Single line text",
        "Rand value; stored as text per live CSV schema (values like 'R150000.00' or '2.5e+007')"])
    r = replace_in_table(t, "Programme", ["Programme", "Single line text",
        "⚠ Text, not a Lookup — matches Directorates.Programme"])
    insert_row_after(t, r, ["Directorate (+ Directorate: Title)", "Lookup (Directorates)",
        "Directorate code; 'Directorate: Title' projection shows the full directorate name"])
    replace_in_table(t, "FinancialYear", ["FinancialYear", "Choice", "2026/27 … 2022/23, colour-coded"])
    changes.append("Projects: ProjectCode + ProjectGoalDescription + Directorate(Lookup), Programme -> text, Budget -> text")

    # --- 4.3 Activities (table 12) ---
    t = d.tables[12]
    r = replace_in_table(t, "Title", ["Title", "Single line text", "Activity ID, e.g. 'ICTM-AC-0001'"])
    r1 = insert_row_after(t, r, ["ActivityCode", "Single line text", "Activity code (e.g. 'ICTM-AC-0001'); also carried in Title"])
    r2 = insert_row_after(t, r1, ["ActivityShortDescription", "Single line text", "Denormalized short label for gallery display"])
    r3 = insert_row_after(t, r2, ["Directorate", "Single line text", "Directorate short code, e.g. 'ICT'"])
    r4 = insert_row_after(t, r3, ["DirectorateLabel", "Single line text", "Full directorate name (denormalized)"])
    insert_row_after(t, r4, ["ProgrammeLabel", "Single line text", "Programme name (denormalized)"])
    r = replace_in_table(t, "Project", ["Project", "Lookup (Projects)", "Indexed"])
    insert_row_after(t, r, ["Project: Reference", "Lookup projection", "Read-only projected field"])
    r = replace_in_table(t, "Owner", ["Owner", "Person", ""])
    insert_row_after(t, r, ["ActivityOwner", "Person", "Contributor accountable for capture (separate 'ActivityOwner' column)"])
    replace_in_table(t, "DueDayOfMonth", ["DueDate", "Date only", "Report due date (EndDate + 3 days in current data)"])
    r = replace_in_table(t, "Active", ["Active", "Yes/No", ""])
    insert_row_after(t, r, ["Status", "Choice", "Active / Completed / Not Started"])
    changes.append("Activities: added ActivityCode/ActivityShortDescription/Directorate/Labels/Project:Reference/ActivityOwner/Status, DueDayOfMonth -> DueDate")

    # --- 4.4 MonthlyReports (table 13) ---
    t = d.tables[13]
    r = replace_in_table(t, "Title", ["Title", "Single line text", "Report ID, e.g. 'ICTM-RP-0001'"])
    r = insert_row_after(t, r, ["Programme / ProgrammeLabel", "Lookup / Text", "Denormalized programme name"])
    r = insert_row_after(t, r, ["Directorate / DirectorateLabel", "Lookup / Text",
        "⚠ DirectorateLabel stores the short CODE (e.g. 'ICTM')"])
    r = insert_row_after(t, r, ["Project / ProjectOwnerLabel", "Lookup / Text",
        "ProjectOwnerLabel is the denormalized Responsibility person's display name"])
    r = insert_row_after(t, r, ["Activity / ActivityOwnerLabel", "Lookup / Text", ""])
    insert_row_after(t, r, ["Activity: ActivityShortDescription / StartDate / EndDate / DueDate",
        "Lookup projections", "Read-only projected fields"])
    replace_in_table(t, "Quarter", ["Quarter", "Choice", "Q1–Q4 (+ All)"])
    replace_in_table(t, "PlannedActivity", ["PlannedAct", "Multiple lines of text", "Internal name: PlannedActivity"])
    r = replace_in_table(t, "ProgressNarrative", ["isFlagged", "Yes/No", "Manually settable by any reviewer role"])
    r = insert_row_after(t, r, ["ReasonForDeviation", "Multiple lines of text", "Free-text reason when progress deviates from plan"])
    r = insert_row_after(t, r, ["RemedialAction", "Multiple lines of text", "Corrective action taken/planned"])
    insert_row_after(t, r, ["Output", "Multiple lines of text", "Reported output / achievement narrative"])
    changes.append("MonthlyReports: aligned with CSV (labels, projections, PlannedAct, isFlagged, ReasonForDeviation, RemedialAction, Output)")

    # --- 4.10 Users -> APP_Users (table 19) ---
    t = d.tables[19]
    r = replace_in_table(t, "Title", ["Title", "Single line text", "Display name"])
    r = insert_row_after(t, r, ["UserAccount", "Person", "Used in LookUp(APP_Users, UserAccount.Email = User().Email)"])
    replace_in_table(t, "Role", ["Role", "Choice",
        "Administrator / ProgrammeManager / Supervisor / Contributor / DeputyDirectorME"])
    r = replace_in_table(t, "Programme", ["Programme", "Single line text",
        "Scope for Programme Manager / Supervisor (text, matches Programmes.Title)"])
    insert_row_after(t, r, ["Directorate (+ Directorate: DirectorateID)", "Lookup (Directorates)",
        "Directorate code; 'Directorate: DirectorateID' projection"])
    changes.append("APP_Users: Title/UserAccount/Role(+DeputyDirectorME)/Programme(text)/Directorate(Lookup)")

    # --- headings / prose ---
    n = replace_in_doc(d, "4.10 List: Users", "4.10 List: APP_Users")
    changes.append(f"Heading: 4.10 List: Users -> APP_Users ({n} replacements)")
    n = replace_in_doc(d, "Users (Programme lookup) -------> (1) Programmes [scopes a Supervisor/Programme Manager]",
        "APP_Users (Directorate lookup) -------> (1) Directorates [scopes Supervisor/Programme Manager]")
    changes.append(f"ERD: APP_Users lookup chain corrected ({n} replacements)")
    n = replace_in_doc(d, "PlannedActivity", "PlannedAct")
    changes.append(f"Arch prose/code: PlannedActivity -> PlannedAct ({n} replacements)")
    n = replace_in_doc(d, "ProgressNarrative", "Output")
    changes.append(f"Arch prose/code: ProgressNarrative -> Output ({n} replacements)")

    d.save(path)
    return changes


# ------------------------------------------------------------- Phase 1 guide
def fix_phase1(path):
    d = docx.Document(path)
    changes = []

    # TABLE 1: list inventory — ReferenceData -> Directorates
    for row in d.tables[1].rows:
        if row.cells[0].text.strip() == "SharePoint Lists":
            set_cell(row.cells[2],
                "Programmes, Projects, Activities, MonthlyReports, ReportComments, KPIDefinitions, APP_Users, Notifications, AuditLog, Directorates")
            changes.append("Inventory: ReferenceData -> Directorates")
        if row.cells[0].text.strip() == "Lookup relationships":
            set_cell(row.cells[2],
                "Projects.Directorate -> Directorates; Activities.Project -> Projects; MonthlyReports.Activity -> Activities; APP_Users.Directorate -> Directorates; Comments, KPIs, Evidence")
            changes.append("Inventory: lookup chain corrected")

    # 5.1 Site Contents checkbox list: ReferenceData -> Directorates
    n = replace_in_doc(d, "ReferenceData", "Directorates")
    changes.append(f"5.1 checkbox list: ReferenceData -> Directorates ({n} replacements)")

    # 5.2 spot-check table (TABLE 11)
    t = d.tables[11]
    replace_in_table(t, "Programmes", ["Programmes", "ProgrammeCode, ProgrammeManager, Active"])
    replace_in_table(t, "Projects", ["Projects", "ProjectID (Title), ProjectCode, KeyDeliverable, FinancialYear, Status, Directorate (Lookup)"])
    replace_in_table(t, "Activities", ["Activities", "ActivityCode, Frequency, ActivityOwner, Supervisor, DueDate, Status"])
    replace_in_table(t, "MonthlyReports", ["MonthlyReports", "Activity (Lookup), ReportingMonth, Status, PercentageComplete, IsOverdue, isFlagged"])
    changes.append("5.2 spot-check columns corrected to live CSV schema")
    # add Directorates row after Programmes
    r = None
    for row in t.rows:
        if row.cells[0].text.strip() == "Programmes":
            r = row
            break
    insert_row_after(t, r, ["Directorates", "DirectorateID (Title), DirectorateName, Programme, Director, DeputyDirector, Deligate, Active"])
    changes.append("5.2 spot-check: added Directorates row")

    # 5.3 lookup verification paragraph
    n = replace_in_doc(d,
        "Open the Projects list. In List Settings → Columns, click 'Programme'. Confirm the Lookup list shows 'Programmes'. Repeat for Activities → Project, and MonthlyReports → Activity.",
        "Open the Projects list. In List Settings → Columns, click 'Directorate'. Confirm the Lookup list shows 'Directorates'. Repeat for Activities → Project (lookup to Projects), and MonthlyReports → Activity (lookup to Activities). Note: Programme is a plain text column on Projects/Directorates, not a Lookup.")
    changes.append(f"5.3 lookup verification text corrected ({n} replacements)")

    # 5.4 ReferenceData seeded -> choice data note
    n = replace_in_doc(d,
        "Open the ReferenceData list. Confirm it has 8 items covering FinancialYear (2024/25 through 2027/28) and Quarter (Q1–Q4).",
        "FinancialYear and Quarter are Choice columns, not a separate lookup list. Per the CSV exports: MonthlyReports.FinancialYear = 2026/27, 2025/26, 2024/25, All; MonthlyReports.Quarter = Q1–Q4, All; Projects.FinancialYear = 2026/27 … 2022/23.")
    changes.append(f"5.4 choice-data note aligned to CSV ({n} replacements)")
    n = replace_in_doc(d, "5.4 ReferenceData seeded", "5.4 Choice data (FinancialYear / Quarter)")
    changes.append(f"5.4 heading updated ({n} replacements)")

    d.save(path)
    return changes


if __name__ == "__main__":
    print("== BRD/FDS/TDS ==")
    for c in fix_brd("APP-MRMS_BRD_FDS_TDS_v2.docx"):
        print(" -", c)
    print("== Architecture Pack ==")
    for c in fix_arch("APP_MRMS_Architecture_Pack.docx"):
        print(" -", c)
    print("== Phase 1 Execution Guide ==")
    for c in fix_phase1("APPMRMS_Phase1_ExecutionGuide.docx"):
        print(" -", c)
    print("Done.")
